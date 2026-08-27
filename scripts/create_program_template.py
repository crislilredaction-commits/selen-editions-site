from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("public/templates/modele-programme-formation-selen.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for name, size, color, before, after in [
    ("Title", 24, "8A4B24", 0, 12),
    ("Heading 1", 16, "8A4B24", 18, 10),
    ("Heading 2", 13, "8A4B24", 14, 7),
]:
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = True
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

header = section.header.paragraphs[0]
header.text = "SELEN ÉDITIONS · TRAME DE PROGRAMME"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header.runs[0].font.size = Pt(8)
header.runs[0].font.color.rgb = RGBColor.from_string("7A6A5E")

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(12)
title_run = title.add_run("Programme de formation")
title_run.font.name = "Calibri"
title_run.font.size = Pt(24)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor.from_string("8A4B24")
subtitle = doc.add_paragraph("Trame éditable à compléter puis à importer dans Selen Daily")
subtitle.runs[0].italic = True
subtitle.runs[0].font.color.rgb = RGBColor.from_string("6B625C")

def field(label: str, hint: str = ""):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    if hint:
        h = p.add_run(f"  {hint}")
        h.italic = True
        h.font.color.rgb = RGBColor.from_string("6B625C")
    line = doc.add_paragraph("________________________________________________________________________________")
    line.paragraph_format.space_after = Pt(8)
    line.runs[0].font.color.rgb = RGBColor.from_string("C9B6A6")

doc.add_heading("1. Identification de la formation", level=1)
field("Intitulé de la formation")
field("Organisme de formation")
field("Public visé")
field("Durée", "en heures et en jours")
field("Modalité", "présentiel, distanciel ou mixte")
field("Tarif TTC")

doc.add_heading("2. Objectifs", level=1)
field("Objectif principal", "commencer par un verbe d’action à l’infinitif")
field("Objectifs pédagogiques", "décrire plusieurs résultats observables ou évaluables")

doc.add_heading("3. Prérequis", level=1)
doc.add_paragraph("Indiquez uniquement les conditions réellement nécessaires. Tout prérequis déclaré devra être vérifié pour chaque apprenant.")
field("Prérequis")

doc.add_heading("4. Contenu détaillé", level=1)
doc.add_paragraph("Présentez les séquences, thèmes et activités dans leur ordre logique. Vous pouvez ajouter ou supprimer des lignes.")
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
headers = ["Séquence / thème", "Contenu et activités", "Durée indicative"]
for cell, text in zip(table.rows[0].cells, headers):
    cell.text = text
    cell.paragraphs[0].runs[0].bold = True
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F4E9DF")
    cell._tc.get_or_add_tcPr().append(shading)
for _ in range(6):
    cells = table.add_row().cells
    cells[0].text = ""
    cells[1].text = ""
    cells[2].text = ""
table.autofit = False
widths = [Inches(1.55), Inches(4.0), Inches(0.95)]
for row in table.rows:
    for cell, width in zip(row.cells, widths):
        cell.width = width

doc.add_heading("5. Moyens pédagogiques et techniques", level=1)
doc.add_paragraph("Décrivez les supports, outils, matériels, méthodes pédagogiques et l’alternance entre théorie et pratique.")
field("Moyens mobilisés")

doc.add_heading("6. Modalités d’évaluation", level=1)
field("Évaluation des acquis", "questionnaire, exercice, étude de cas, mise en situation…")
field("Positionnement avant la formation", "questionnaire de connaissances et modalités de réalisation")

doc.add_heading("7. Accès et contacts", level=1)
field("Délai d’accès")
field("Téléphone de l’organisme")
field("Email de l’organisme")
field("Site Internet de l’organisme", "facultatif")

footer = section.footer.paragraphs[0]
footer.text = "Trame Selen Daily · À adapter à la formation proposée"
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.size = Pt(8)
footer.runs[0].font.color.rgb = RGBColor.from_string("7A6A5E")

doc.save(OUT)
print(OUT)
